import os
import io
import json
import base64
import logging
import re
import uuid
import time
from datetime import datetime, timezone, timedelta
from typing import Optional, List
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Request, status, Query
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr, Field
from passlib.context import CryptContext
import jwt
import certifi
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo.errors import DuplicateKeyError
from dotenv import load_dotenv
from openai import OpenAI
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# Native AI & Processing Libraries
import PyPDF2
import docx
import easyocr
import faiss
from PIL import Image
from sentence_transformers import SentenceTransformer

# Startup Logging Configuration
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

# Import shared services (database, LLM client, JSON extractor)
from services import db, groq_client, extract_json

# --- Rate Limiter Setup (Issue 3) ---
limiter = Limiter(key_func=get_remote_address)

# --- Modern Lifespan Event Handler (Issue 4 from previous & DB indexes) ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        await db.users.create_index("email", unique=True)
        logger.info("MongoDB unique index on 'email' verified successfully.")
    except Exception as e:
        logger.error(f"Failed to create unique index on email: {e}")

    try:
        await db.sessions.create_index("session_id", unique=True)
        logger.info("MongoDB unique index on 'session_id' verified successfully.")
    except Exception as e:
        logger.error(f"Failed to create unique index on session_id: {e}")
    
    logger.info("Startup complete. LearnNova AI Production Backend is ready.")
    yield

app = FastAPI(
    title="LearnNova AI Enterprise Production Backend",
    lifespan=lifespan
)

app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# CORS Configuration
allow_origins_list = os.getenv("ALLOWED_ORIGINS", "*").split(",") if os.getenv("ALLOWED_ORIGINS") else ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins_list,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- MongoDB Setup ---
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")
DB_NAME = os.getenv("DB_NAME", "Lnnova_DB")

client_kwargs = {"serverSelectionTimeoutMS": 5000}
if MONGO_URI.startswith("mongodb+srv://"):
    client_kwargs["tls"] = True
    client_kwargs["tlsCAFile"] = certifi.where()

client = AsyncIOMotorClient(MONGO_URI, **client_kwargs)
db = client[DB_NAME]

# --- Auth Security Config ---
JWT_SECRET = os.getenv("JWT_SECRET")
if not JWT_SECRET:
    raise RuntimeError("JWT_SECRET is missing in .env")
ALGORITHM = os.getenv("ALGORITHM", "HS256")
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise RuntimeError("GROQ_API_KEY not found in .env")

groq_client = OpenAI(
    base_url="https://api.groq.com/openai/v1",
    api_key=GROQ_API_KEY
)

# --- Lazy Loading Embedding Model (prevents OOM on startup in low-memory environments) ---
_embedding_model = None
def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        logger.info("Lazy loading SentenceTransformer model...")
        _embedding_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
    return _embedding_model

# --- Lazy Loading OCR Reader (Issue 7) ---
_ocr_reader = None
def get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        logger.info("Lazy loading EasyOCR reader...")
        _ocr_reader = easyocr.Reader(['en', 'hi'], gpu=False)
    return _ocr_reader

# Local Disk or Cloud Storage Fallback Dir (Issue 1)
INDEX_DIR = "./faiss_indices"
os.makedirs(INDEX_DIR, exist_ok=True)

vector_indices = {}  
index_chunks = {}    
# Note: document_texts removed to prevent high RAM usage on large PDFs (Issue 13)

# --- PYDANTIC MODELS ---
class UserRegister(BaseModel):
    name: str
    email: EmailStr
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user_name: str
    email: str

class ChatQuery(BaseModel):
    message: str
    language: str = "English"
    session_id: Optional[str] = None  # Issue 17: Auto-generated UUID if none
    doc_name: Optional[str] = None

class ChatResponse(BaseModel):
    response: str
    session_id: str

class QuizAttemptSubmit(BaseModel):
    score: int
    total_questions: int

# --- HELPER & SECURITY FUNCTIONS ---
def hash_password(password: str) -> str:
    safe_pwd = password.encode('utf-8')[:72].decode('utf-8', errors='ignore')
    return pwd_context.hash(safe_pwd)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    safe_pwd = plain_password.encode('utf-8')[:72].decode('utf-8', errors='ignore')
    return pwd_context.verify(safe_pwd, hashed_password)

def create_access_token(data: dict) -> str:
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, JWT_SECRET, algorithm=ALGORITHM)

# Issue 2 & 19: JWT Dependency & Verification
async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> str:
    token = credentials.credentials
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid authentication credentials.")
        
        # Check if token/user session is blacklisted (Logout check - Issue 19)
        blacklisted = await db.blacklist.find_one({"token": token})
        if blacklisted:
            raise HTTPException(status_code=401, detail="Token has been invalidated (Logged out).")
            
        return user_id
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Could not validate credentials.")

# Issue 11: Strong Password Policy
def validate_password_strength(password: str):
    if len(password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters long.")
    if not re.search(r"[A-Z]", password):
        raise HTTPException(status_code=400, detail="Password must contain at least one uppercase letter.")
    if not re.search(r"[a-z]", password):
        raise HTTPException(status_code=400, detail="Password must contain at least one lowercase letter.")
    if not re.search(r"\d", password):
        raise HTTPException(status_code=400, detail="Password must contain at least one digit.")
    if not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):
        raise HTTPException(status_code=400, detail="Password must contain at least one special character.")

def split_text_into_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += (chunk_size - overlap)
    return chunks

# Issue 6: Robust Regex JSON Extraction
def extract_json(content: str):
    content = content.strip()
    try:
        return json.loads(content)
    except Exception:
        pass

    # Try Regex for lists or dicts
    for pattern in [r'\[.*\]', r'\{.*\}']:
        match = re.search(pattern, content, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                continue

    raise HTTPException(status_code=500, detail="Failed to parse structured JSON response from AI.")

# --- AUTH ENDPOINTS ---
@app.post("/api/auth/register", response_model=TokenResponse)
@limiter.limit("5/minute")
async def register(request: Request, user: UserRegister):
    validate_password_strength(user.password)
    try:
        user_email = user.email.lower().strip()
        user_doc = {
            "name": user.name.strip(),
            "email": user_email,
            "password": hash_password(user.password),
            "created_at": datetime.now(timezone.utc)
        }
        result = await db.users.insert_one(user_doc)
        token = create_access_token({"sub": str(result.inserted_id), "email": user_email})
        return {"access_token": token, "token_type": "bearer", "user_name": user.name, "email": user_email}
    except DuplicateKeyError:
        raise HTTPException(status_code=400, detail="User with this email already exists.")
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Internal server error during registration.")

# Refreshed bcrypt version compatibility check
@app.post("/api/auth/login", response_model=TokenResponse)
@limiter.limit("10/minute")
async def login(request: Request, user: UserLogin):
    try:
        user_email = user.email.lower().strip()
        db_user = await db.users.find_one({"email": user_email})
        if not db_user or not verify_password(user.password, db_user["password"]):
            raise HTTPException(status_code=400, detail="Invalid email address or password.")

        token = create_access_token({"sub": str(db_user["_id"]), "email": db_user["email"]})
        return {"access_token": token, "token_type": "bearer", "user_name": db_user["name"], "email": db_user["email"]}
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Internal server error during login.")

# Issue 19: Logout Endpoint
@app.post("/api/auth/logout")
async def logout(credentials: HTTPAuthorizationCredentials = Depends(security), user_id: str = Depends(get_current_user)):
    token = credentials.credentials
    try:
        await db.blacklist.insert_one({"token": token, "blacklisted_at": datetime.now(timezone.utc)})
        return {"status": "success", "message": "Successfully logged out."}
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Logout failed.")

# --- DOCUMENT MANAGEMENT & UPLOAD (Issues 5, 8, 12, 13, 20) ---
@app.post("/api/upload")
@limiter.limit("10/minute")
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    doc_name: str = Form(...),
    user_id: str = Depends(get_current_user)
):
    file_bytes = await file.read()
    MAX_UPLOAD = 15 * 1024 * 1024  # 15 MB
    if len(file_bytes) > MAX_UPLOAD:
        raise HTTPException(status_code=413, detail="File too large (Max 15MB).")

    # Issue 12: Strict MIME & Extension Validation
    allowed_mimes = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/webp": ".webp",
        "text/plain": ".txt"
    }
    
    mime_type = file.content_type
    if mime_type not in allowed_mimes:
        raise HTTPException(status_code=400, detail=f"Unsupported MIME type: {mime_type}. Supported: PDF, DOCX, TXT, PNG, JPG, WEBP.")

    extracted_text = ""
    try:
        if mime_type == "application/pdf":
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            except Exception:
                raise HTTPException(status_code=400, detail="Invalid or corrupted PDF file.")
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"

        elif mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    if para.text.strip():
                        extracted_text += para.text + "\n"
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            extracted_text += row_text + "\n"
            except Exception:
                raise HTTPException(status_code=400, detail="Invalid or corrupted DOCX file.")
                    
        elif mime_type.startswith("image/"):
            try:
                img = Image.open(io.BytesIO(file_bytes))
                # Resize large images for ultra-fast CPU OCR processing
                max_dim = 1280
                if max(img.size) > max_dim:
                    img.thumbnail((max_dim, max_dim))
                    buffer = io.BytesIO()
                    img.save(buffer, format="JPEG", quality=85)
                    file_bytes_ocr = buffer.getvalue()
                else:
                    file_bytes_ocr = file_bytes

                reader = get_ocr_reader()
                ocr_results = reader.readtext(file_bytes_ocr, detail=0)
                ocr_text = " ".join(ocr_results).strip()
                extracted_text = f"Extracted Text from Image ({file.filename}): {ocr_text}" if ocr_text else f"Image document {file.filename} uploaded."
            except Exception as img_err:
                logger.warning(f"Fast OCR processing failed: {img_err}")
                extracted_text = f"Image document {file.filename} uploaded."
        else:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract any readable text from the file.")

        # High-performance chunking (chunk_size=1000 for 50% fewer embeddings and faster indexing)
        chunks = split_text_into_chunks(extracted_text, chunk_size=1000, overlap=100)
        if len(chunks) == 0:
            raise HTTPException(status_code=400, detail="No chunks generated.")

        # Batch encode embeddings with 64 batch size for maximum speed
        embedding_model = get_embedding_model()
        embeddings = embedding_model.encode(
            chunks,
            batch_size=64,
            convert_to_numpy=True,
            show_progress_bar=False
        ).astype('float32')
        
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings)

        # Issue 5: Unique sub-folder per user to prevent overwrites
        clean_doc_name = re.sub(r'[^a-zA-Z0-9_-]', '_', doc_name.lower())
        user_specific_dir = os.path.join(INDEX_DIR, user_id)
        os.makedirs(user_specific_dir, exist_ok=True)

        index_path = os.path.join(user_specific_dir, f"{clean_doc_name}.index")
        chunks_path = os.path.join(user_specific_dir, f"{clean_doc_name}_chunks.json")

        faiss.write_index(index, index_path)
        with open(chunks_path, "w", encoding="utf-8") as f:
            json.dump(chunks, f)

        file_ext = file.filename.split('.')[-1].lower() if '.' in file.filename else 'txt'
        file_size = len(file_bytes)

        # Issue 13: Store metadata record in MongoDB with size, type, and favorite status
        await db.user_documents.update_one(
            {"user_id": user_id, "doc_name": clean_doc_name},
            {
                "$set": {
                    "user_id": user_id,
                    "doc_name": clean_doc_name,
                    "original_filename": file.filename,
                    "file_size": file_size,
                    "file_type": file_ext,
                    "chunks_count": len(chunks),
                    "index_path": index_path,
                    "chunks_path": chunks_path,
                    "updated_at": datetime.now(timezone.utc)
                },
                "$setOnInsert": {
                    "is_favorite": False,
                    "created_at": datetime.now(timezone.utc)
                }
            },
            upsert=True
        )

        return {
            "status": "success",
            "filename": file.filename,
            "document_key": clean_doc_name,
            "chunks_indexed": len(chunks),
            "file_size": file_size,
            "file_type": file_ext,
            "message": f"Successfully uploaded and indexed '{file.filename}'!"
        }
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Error processing and uploading file.")

# Issue 20: Delete Document API
@app.delete("/api/document/{doc_name}")
async def delete_document(doc_name: str, user_id: str = Depends(get_current_user)):
    try:
        doc_record = await db.user_documents.find_one({"user_id": user_id, "doc_name": doc_name})
        if not doc_record:
            raise HTTPException(status_code=404, detail="Document not found.")

        # Remove local files
        if os.path.exists(doc_record["index_path"]):
            os.remove(doc_record["index_path"])
        if os.path.exists(doc_record["chunks_path"]):
            os.remove(doc_record["chunks_path"])

        # Remove from MongoDB record
        await db.user_documents.delete_one({"user_id": user_id, "doc_name": doc_name})
        return {"status": "success", "message": f"Document '{doc_name}' deleted successfully."}
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to delete document.")

# Get All User Documents API for My Library
@app.get("/api/documents")
async def get_user_documents(user_id: str = Depends(get_current_user)):
    try:
        docs = await db.user_documents.find({"user_id": user_id}).sort("updated_at", -1).to_list(length=100)
        formatted_docs = []
        for d in docs:
            formatted_docs.append({
                "doc_name": d.get("doc_name"),
                "original_filename": d.get("original_filename", d.get("doc_name")),
                "chunks_count": d.get("chunks_count", 0),
                "file_size": d.get("file_size", d.get("chunks_count", 0) * 1000),
                "file_type": d.get("file_type", d.get("original_filename", "").split(".")[-1].lower() if "." in d.get("original_filename", "") else "pdf"),
                "is_favorite": d.get("is_favorite", False),
                "created_at": d.get("created_at").isoformat() if d.get("created_at") else (d.get("updated_at").isoformat() if d.get("updated_at") else None),
                "updated_at": d.get("updated_at").isoformat() if d.get("updated_at") else None
            })
        return {"status": "success", "documents": formatted_docs}
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to fetch user documents.")

# --- 📚 MY LIBRARY EXTENDED API ENDPOINTS ---
class RenameDocRequest(BaseModel):
    new_title: str

@app.put("/api/document/{doc_name}/rename")
async def rename_document(doc_name: str, payload: RenameDocRequest, user_id: str = Depends(get_current_user)):
    try:
        new_title = payload.new_title.strip()
        if not new_title:
            raise HTTPException(status_code=400, detail="New title cannot be empty.")
        
        result = await db.user_documents.update_one(
            {"user_id": user_id, "doc_name": doc_name},
            {"$set": {"original_filename": new_title, "updated_at": datetime.now(timezone.utc)}}
        )
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="Document not found.")
        return {"status": "success", "message": f"Document renamed to '{new_title}'."}
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to rename document.")

@app.post("/api/document/{doc_name}/favorite")
async def toggle_favorite(doc_name: str, user_id: str = Depends(get_current_user)):
    try:
        doc_record = await db.user_documents.find_one({"user_id": user_id, "doc_name": doc_name})
        if not doc_record:
            raise HTTPException(status_code=404, detail="Document not found.")
        new_fav_status = not doc_record.get("is_favorite", False)
        await db.user_documents.update_one(
            {"user_id": user_id, "doc_name": doc_name},
            {"$set": {"is_favorite": new_fav_status, "updated_at": datetime.now(timezone.utc)}}
        )
        return {"status": "success", "is_favorite": new_fav_status}
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to toggle favorite status.")

@app.get("/api/document/{doc_name}/download")
async def download_document(doc_name: str, user_id: str = Depends(get_current_user)):
    try:
        doc_record = await db.user_documents.find_one({"user_id": user_id, "doc_name": doc_name})
        if not doc_record:
            raise HTTPException(status_code=404, detail="Document not found.")
        chk_path = doc_record.get("chunks_path")
        if not chk_path or not os.path.exists(chk_path):
            raise HTTPException(status_code=404, detail="Document content chunks not found.")
        with open(chk_path, "r", encoding="utf-8") as f:
            chunks = json.load(f)
        full_text = "\n\n".join(chunks)
        filename = doc_record.get("original_filename", f"{doc_name}.txt")
        return {"status": "success", "filename": filename, "content": full_text}
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to download document content.")

@app.post("/api/document/{doc_name}/summarize")
@limiter.limit("10/minute")
async def summarize_document(request: Request, doc_name: str, user_id: str = Depends(get_current_user)):
    try:
        doc_record = await db.user_documents.find_one({"user_id": user_id, "doc_name": doc_name})
        if not doc_record:
            raise HTTPException(status_code=404, detail="Document not found.")
        chk_path = doc_record.get("chunks_path")
        if not chk_path or not os.path.exists(chk_path):
            raise HTTPException(status_code=400, detail="No readable text available for this document.")
        with open(chk_path, "r", encoding="utf-8") as f:
            chunks = json.load(f)
        # Use first 20 chunks (~20k chars) for summarization
        context = "\n---\n".join(chunks[:20])
        if not context.strip():
            raise HTTPException(status_code=400, detail="No readable text available for this document.")
        
        prompt = f"""Summarize the following document content professionally in Markdown.
Include:
1. **Executive Summary** (2-3 sentences)
2. **Core Takeaways & Key Concepts** (Bullet points)
3. **Main Terminology & Definitions**

Document Content:
{context}"""
        res = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        summary_text = res.choices[0].message.content
        return {"status": "success", "doc_name": doc_name, "summary": summary_text}
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to generate document summary.")

# --- HELPER: RETRIEVE RELEVANT CHUNKS ACROSS USER DOCS (Issues 14, 15, 16) ---
async def get_relevant_context_for_user(user_id: str, query_text: str, k_per_doc: int = 4, doc_name: Optional[str] = None) -> str:
    query_filter = {"user_id": user_id}
    if doc_name:
        query_filter["doc_name"] = doc_name

    user_docs = await db.user_documents.find(query_filter).to_list(length=20)
    if not user_docs and doc_name:
        user_docs = await db.user_documents.find({"user_id": user_id}).to_list(length=20)
    if not user_docs:
        return ""

    all_retrieved_chunks = []
    query_embedding = get_embedding_model().encode([query_text], convert_to_numpy=True).astype('float32')

    for doc in user_docs:
        idx_path = doc["index_path"]
        chk_path = doc["chunks_path"]
        if os.path.exists(idx_path) and os.path.exists(chk_path):
            try:
                index = faiss.read_index(idx_path)
                with open(chk_path, "r", encoding="utf-8") as f:
                    chunks = json.load(f)
                
                k = min(k_per_doc, len(chunks))
                distances, indices = index.search(query_embedding, k)
                for idx in indices[0]:
                    if idx != -1 and idx < len(chunks):
                        all_retrieved_chunks.append(chunks[idx])
            except Exception as e:
                logger.error(f"Error reading index for doc {doc['doc_name']}: {e}")

    return "\n---\n".join(all_retrieved_chunks)

# --- VISION ANALYZE ENDPOINT ---
@app.post("/api/vision-analyze")
@limiter.limit("15/minute")
async def vision_analyze(
    request: Request,
    prompt: str = Form(...), 
    language: str = Form("English"),
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user)
):
    try:
        image_bytes = await file.read()
        if len(image_bytes) > 10 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="Image file too large.")

        mime_type = file.content_type or "image/jpeg"
        base64_image = base64.b64encode(image_bytes).decode('utf-8')

        # Fetch relevant context using semantic search
        context = await get_relevant_context_for_user(user_id, prompt, k_per_doc=3)
        system_note = f"You are an expert AI Education Assistant. Answer clearly and professionally in {language}."
        context_note = f"Document Context: {context}\n\n" if context else ""
        full_prompt = f"{system_note}\n\n{context_note}User Question: {prompt}"

        # Try qwen/qwen3.6-27b (vision-capable) with up to 3 retries on 503 over-capacity errors
        vision_response = None
        last_error = None
        for attempt in range(3):
            try:
                vision_completion = groq_client.chat.completions.create(
                    model="qwen/qwen3.6-27b",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": full_prompt},
                                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{base64_image}"}}
                            ]
                        }
                    ]
                )
                vision_response = vision_completion.choices[0].message.content
                break  # Success — exit retry loop
            except Exception as ex:
                last_error = ex
                err_str = str(ex)
                if "503" in err_str or "over capacity" in err_str.lower():
                    logger.warning(f"Vision model over capacity (attempt {attempt+1}/3). Retrying in {2 ** attempt}s...")
                    time.sleep(2 ** attempt)  # 1s, 2s, 4s backoff
                else:
                    break  # Non-retryable error, exit immediately

        if vision_response:
            return {"analysis": vision_response}

        # Fallback: OCR + text-only LLM analysis
        logger.warning(f"Vision model unavailable ({last_error}). Falling back to OCR + text analysis.")
        try:
            reader = get_ocr_reader()
            ocr_results = reader.readtext(image_bytes, detail=0)
            ocr_text = " ".join(ocr_results).strip()
        except Exception:
            ocr_text = ""

        fallback_text = f"Image OCR text: {ocr_text}" if ocr_text else "No readable text found in the image."
        fallback_prompt = f"{system_note}\n\n{context_note}The user uploaded an image. Based on the extracted text below, answer: {prompt}\n\n{fallback_text}"

        text_completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": fallback_prompt}],
            temperature=0.3,
            max_tokens=1024
        )
        fallback_note = "\n\n*(Note: Direct image AI is temporarily over capacity. Analysis based on text extracted from the image.)*"
        return {"analysis": text_completion.choices[0].message.content + fallback_note}

    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail=f"Vision analysis failed: {str(e)}")

# --- MONGODB SESSION STORAGE HELPERS (Issue 9: Extended Context 25 messages) ---
async def get_session_history(session_id: str) -> list:
    session_doc = await db.sessions.find_one({"session_id": session_id})
    if session_doc and "messages" in session_doc:
        return session_doc["messages"]
    return []

async def save_session_history(session_id: str, messages: list):
    trimmed_messages = messages[-25:]  # Issue 9: Kept 25 messages for better continuity
    await db.sessions.update_one(
        {"session_id": session_id},
        {"$set": {"messages": trimmed_messages, "updated_at": datetime.now(timezone.utc)}},
        upsert=True
    )

# --- SECURE CHAT ENDPOINT (Issues 2, 3, 9, 17) ---
async def process_chat(query: ChatQuery, user_id: str):
    try:
        # Issue 17: Generate secure session UUID if not provided
        session_id = query.session_id or str(uuid.uuid4())
        
        # Retrieve relevant context from user FAISS indexes (with doc_name filter if specified)
        context = await get_relevant_context_for_user(user_id, query.message, k_per_doc=4, doc_name=query.doc_name)

        session_messages = await get_session_history(session_id)
        session_messages.append({"role": "user", "content": query.message})

        system_content = f"""You are LearnNova AI, an intelligent AI Smart Education Assistant and Personal Tutor.

PRIMARY INSTRUCTIONS:
- If reference context from student's uploaded document is provided below, PRIORITIZE answering the user's question directly using that document context.
- If the user asks a question about the document (e.g. subject code, title, topics, definitions, syllabus details), extract the exact details from the document context.
- For general questions or greetings without document context, answer politely as a helpful tutor.

STRICT FORMATTING & LANGUAGE RULES:
- Respond fluently in the requested language: {query.language}.
- Bold key terms, codes, and definitions for clarity.
- Use clean Markdown lists and professional structure.

Document Reference Context:
{context if context.strip() else "No document uploaded yet or no relevant document context available."}
"""

        messages = [{"role": "system", "content": system_content}] + session_messages

        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.3,
            max_tokens=1024,
        )

        response_text = completion.choices[0].message.content
        session_messages.append({"role": "assistant", "content": response_text})
        
        await save_session_history(session_id, session_messages)
        return ChatResponse(response=response_text, session_id=session_id)
    except HTTPException as h:
        raise h
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Internal chat generation error.")

@app.post("/api/chat", response_model=ChatResponse)
@limiter.limit("30/minute")
async def chat_endpoint(request: Request, query: ChatQuery, user_id: str = Depends(get_current_user)):
    return await process_chat(query, user_id)

# --- HELPER: GET FULL CONTEXT FROM SPECIFIC OR LATEST UPLOADED DOCUMENT ---
async def get_document_context_for_user(user_id: str, max_chunks: int = 15, doc_name: Optional[str] = None) -> str:
    """Fetch text chunks from a specific document or the user's most recently uploaded document."""
    query_filter = {"user_id": user_id}
    if doc_name:
        query_filter["doc_name"] = doc_name

    target_doc = await db.user_documents.find_one(
        query_filter,
        sort=[("updated_at", -1)]
    )
    if not target_doc and doc_name:
        target_doc = await db.user_documents.find_one(
            {"user_id": user_id},
            sort=[("updated_at", -1)]
        )

    if not target_doc:
        return ""

    chk_path = target_doc.get("chunks_path")
    if chk_path and os.path.exists(chk_path):
        try:
            with open(chk_path, "r", encoding="utf-8") as f:
                chunks = json.load(f)
            selected_chunks = chunks[:max_chunks]
            return "\n---\n".join(selected_chunks)
        except Exception as e:
            logger.error(f"Error reading chunks for doc {target_doc.get('doc_name')}: {e}")

    return ""

# --- SMART GENERATION ENDPOINTS (Issues 14, 15, 16) ---
@app.post("/api/generate-quiz")
@limiter.limit("5/minute")
async def generate_quiz(request: Request, doc_name: Optional[str] = Query(None), user_id: str = Depends(get_current_user)):
    context = await get_document_context_for_user(user_id, max_chunks=15, doc_name=doc_name)
    if not context.strip():
        context = await get_relevant_context_for_user(user_id, "important definitions core concepts summary exam", k_per_doc=5, doc_name=doc_name)
    if not context.strip():
        raise HTTPException(status_code=400, detail="Please upload a document first.")

    prompt = f"""Based STRICTLY on the following uploaded document content, create 10 high-quality multiple-choice questions in strict JSON format.
Document Content:
{context}

Return ONLY valid JSON structure:
[
  {{
    "id": 1,
    "question": "Question text?",
    "options": ["Option A", "Option B", "Option C", "Option D"],
    "correctAnswer": 0,
    "explanation": "Explanation."
  }}
]"""
    res = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return {"quizzes": extract_json(res.choices[0].message.content)}

@app.post("/api/generate-flashcards")
@limiter.limit("5/minute")
async def generate_flashcards(request: Request, doc_name: Optional[str] = Query(None), user_id: str = Depends(get_current_user)):
    context = await get_document_context_for_user(user_id, max_chunks=15, doc_name=doc_name)
    if not context.strip():
        context = await get_relevant_context_for_user(user_id, "key terms definitions concepts formulas", k_per_doc=5, doc_name=doc_name)
    if not context.strip():
        raise HTTPException(status_code=400, detail="Please upload a document first.")

    prompt = f"""Extract 10 key terms and definitions as flashcards STRICTLY from the following uploaded document content in strict JSON format.
Document Content:
{context}

Return ONLY valid JSON structure:
[
  {{
    "id": 1,
    "front": "Term",
    "back": "Definition"
  }}
]"""
    res = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return {"flashcards": extract_json(res.choices[0].message.content)}

@app.post("/api/generate-study-plan")
@limiter.limit("5/minute")
async def generate_study_plan(
    request: Request,
    language: str = Form("English"), 
    start_date: str = Form(...), 
    exam_date: str = Form(...),
    doc_name: Optional[str] = Form(None),
    user_id: str = Depends(get_current_user)
):
    try:
        d1 = datetime.strptime(start_date, "%Y-%m-%d")
        d2 = datetime.strptime(exam_date, "%Y-%m-%d")
        total_days = (d2 - d1).days + 1
        if total_days <= 0:
            raise HTTPException(status_code=400, detail="Exam date must be after start date.")
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format (Use YYYY-MM-DD).")

    context = await get_document_context_for_user(user_id, max_chunks=15, doc_name=doc_name)
    if not context.strip():
        context = await get_relevant_context_for_user(user_id, "syllabus topics modules chapters overview", k_per_doc=6, doc_name=doc_name)
    if not context.strip():
        raise HTTPException(status_code=400, detail="Please upload a document first.")

    prompt = f"""Create a day-by-day structured study plan spanning {total_days} days from {start_date} to {exam_date} in {language} based STRICTLY on the uploaded document content below.
Document Content:
{context}

Return ONLY valid JSON structure:
[
  {{
    "day": 1,
    "date": "{start_date}",
    "title": "Title",
    "tasks": ["Task 1", "Task 2"]
  }}
]"""
    res = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return {"study_plan": extract_json(res.choices[0].message.content)}

# --- REAL-TIME DASHBOARD & ACTIVITY ENDPOINTS ---
async def record_user_activity(user_id: str, minutes: int = 5):
    try:
        today_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        await db.user_activity.update_one(
            {"user_id": user_id, "date": today_str},
            {"$inc": {"actions_count": 1, "study_minutes": minutes}, "$set": {"updated_at": datetime.now(timezone.utc)}},
            upsert=True
        )
    except Exception as e:
        logger.error(f"Failed to record activity for user {user_id}: {e}")

@app.post("/api/quiz/submit")
async def submit_quiz_result(request: Request, attempt: QuizAttemptSubmit, user_id: str = Depends(get_current_user)):
    percentage = round((attempt.score / max(attempt.total_questions, 1)) * 100)
    doc = {
        "user_id": user_id,
        "score": attempt.score,
        "total_questions": attempt.total_questions,
        "percentage": percentage,
        "created_at": datetime.now(timezone.utc)
    }
    await db.quiz_attempts.insert_one(doc)
    await record_user_activity(user_id, minutes=10)
    return {"status": "success", "percentage": percentage}

@app.post("/api/activity/log")
async def log_activity(request: Request, minutes: int = 5, user_id: str = Depends(get_current_user)):
    await record_user_activity(user_id, minutes=minutes)
    return {"status": "success"}

@app.get("/api/dashboard/stats")
async def get_dashboard_stats(user_id: str = Depends(get_current_user)):
    # 1. Total Documents & Topics (Chunks) Mastered
    user_docs = await db.user_documents.find({"user_id": user_id}).to_list(length=100)
    docs_count = len(user_docs)
    topics_mastered = sum(doc.get("chunks_count", 0) for doc in user_docs)

    # 2. Quiz Attempts & Average Score
    quiz_records = await db.quiz_attempts.find({"user_id": user_id}).to_list(length=100)
    if quiz_records:
        avg_score = round(sum(q.get("percentage", 0) for q in quiz_records) / len(quiz_records))
    else:
        avg_score = 0

    # 3. Study Time & Streak Calculation
    activity_records = await db.user_activity.find({"user_id": user_id}).to_list(length=365)
    total_study_minutes = sum(a.get("study_minutes", 0) for a in activity_records)

    # Fallback heuristic if activity just started
    if total_study_minutes == 0 and docs_count > 0:
        total_study_minutes = docs_count * 15

    hours = total_study_minutes // 60
    mins = total_study_minutes % 60
    study_time_formatted = f"{hours}h {mins}m" if hours > 0 else f"{mins}m"

    # Consecutive active days streak
    dates_active = sorted(set(a["date"] for a in activity_records if a.get("date")), reverse=True)
    streak = 0
    today = datetime.now(timezone.utc).date()
    curr = today
    while True:
        curr_str = curr.strftime("%Y-%m-%d")
        if curr_str in dates_active:
            streak += 1
            curr -= timedelta(days=1)
        elif curr == today:
            curr -= timedelta(days=1)
        else:
            break

    if streak == 0 and (docs_count > 0 or len(quiz_records) > 0):
        streak = 1

    # 4. Weekly Activity Breakdown (Mon - Sun)
    today_dt = datetime.now(timezone.utc)
    start_of_week = today_dt - timedelta(days=today_dt.weekday())
    weekly_activity = []
    for i in range(7):
        day_dt = start_of_week + timedelta(days=i)
        day_str = day_dt.strftime("%Y-%m-%d")
        day_act = next((a for a in activity_records if a.get("date") == day_str), None)
        day_hours = round((day_act.get("study_minutes", 0) / 60) if day_act else 0, 1)
        weekly_activity.append(day_hours)

    weekly_minutes = sum(int(h * 60) for h in weekly_activity)
    weekly_goal_percent = min(100, round((weekly_minutes / 900) * 100))  # 15h goal = 900m

    return {
        "study_time": study_time_formatted,
        "total_study_minutes": total_study_minutes,
        "avg_quiz_score": avg_score,
        "current_streak": streak,
        "topics_mastered": max(topics_mastered, docs_count * 5),
        "documents_count": docs_count,
        "quizzes_taken": len(quiz_records),
        "weekly_activity": weekly_activity,
        "weekly_goal_percent": weekly_goal_percent,
        "weekly_hours_completed": round(weekly_minutes / 60, 1)
    }

# --- 📊 1. FULL ANALYTICS & PERFORMANCE ENDPOINT ---
@app.get("/api/analytics")
async def get_full_analytics(user_id: str = Depends(get_current_user)):
    try:
        user_docs = await db.user_documents.find({"user_id": user_id}).to_list(length=100)
        docs_count = len(user_docs)

        quiz_records = await db.quiz_attempts.find({"user_id": user_id}).sort("created_at", -1).to_list(length=100)
        quizzes_taken = len(quiz_records)
        avg_score = round(sum(q.get("percentage", 0) for q in quiz_records) / quizzes_taken) if quizzes_taken > 0 else 0

        activity_records = await db.user_activity.find({"user_id": user_id}).to_list(length=365)
        total_study_minutes = sum(a.get("study_minutes", 0) for a in activity_records)

        hours = total_study_minutes // 60
        mins = total_study_minutes % 60
        study_time_formatted = f"{hours}h {mins}m" if hours > 0 else f"{mins}m"

        has_data = (docs_count > 0) or (quizzes_taken > 0) or (total_study_minutes > 0)

        # Topic mastery breakdown based ONLY on actual user documents and quiz performance
        topic_mastery = []
        mastered_count = 0
        for doc in user_docs:
            doc_name = doc.get("doc_name")
            title = doc.get("original_filename", doc_name)
            chunks = doc.get("chunks_count", 0)

            # Match quizzes attempted for this document if stored, or fallback to user avg
            doc_quizzes = [q for q in quiz_records if q.get("doc_name") == doc_name]
            if doc_quizzes:
                doc_mastery = round(sum(q.get("percentage", 0) for q in doc_quizzes) / len(doc_quizzes))
            elif quizzes_taken > 0:
                doc_mastery = avg_score
            else:
                doc_mastery = 0

            status = "Mastered" if doc_mastery >= 80 else ("In Progress" if doc_mastery >= 40 else "Needs Review")
            if doc_mastery >= 80:
                mastered_count += 1

            topic_mastery.append({
                "doc_name": doc_name,
                "title": title,
                "chunks": chunks,
                "mastery_percent": doc_mastery,
                "status": status
            })

        # Structured Quiz Attempt History Log (Real user attempts only)
        quiz_history = []
        for q in quiz_records[:10]:
            created_dt = q.get("created_at")
            date_str = created_dt.strftime("%b %d, %Y %H:%M") if isinstance(created_dt, datetime) else "Recent"
            quiz_history.append({
                "score": q.get("score", 0),
                "total_questions": q.get("total_questions", 10),
                "percentage": q.get("percentage", 0),
                "date": date_str
            })

        # Weekly Activity Breakdown (Mon - Sun) from MongoDB
        today_dt = datetime.now(timezone.utc)
        start_of_week = today_dt - timedelta(days=today_dt.weekday())
        weekly_activity = []
        for i in range(7):
            day_dt = start_of_week + timedelta(days=i)
            day_str = day_dt.strftime("%Y-%m-%d")
            day_act = next((a for a in activity_records if a.get("date") == day_str), None)
            day_hours = round((day_act.get("study_minutes", 0) / 60) if day_act else 0, 1)
            weekly_activity.append(day_hours)

        retention_rate = round(avg_score * 0.9) if quizzes_taken > 0 else 0

        return {
            "status": "success",
            "has_data": has_data,
            "study_time": study_time_formatted,
            "total_study_minutes": total_study_minutes,
            "avg_quiz_score": avg_score,
            "quizzes_taken": quizzes_taken,
            "documents_count": docs_count,
            "topics_mastered": mastered_count,
            "weekly_activity": weekly_activity,
            "topic_mastery": topic_mastery,
            "quiz_history": quiz_history,
            "retention_rate": retention_rate
        }
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to fetch analytics data.")

# --- 🧬 2. PERSONALIZED LEARNING DNA ENDPOINT ---
@app.get("/api/learning-dna")
async def get_learning_dna(user_id: str = Depends(get_current_user)):
    try:
        user_docs = await db.user_documents.find({"user_id": user_id}).to_list(length=50)
        quiz_records = await db.quiz_attempts.find({"user_id": user_id}).to_list(length=50)
        activity_records = await db.user_activity.find({"user_id": user_id}).to_list(length=100)

        docs_count = len(user_docs)
        quizzes_taken = len(quiz_records)
        total_study_minutes = sum(a.get("study_minutes", 0) for a in activity_records)
        has_data = (docs_count > 0) or (quizzes_taken > 0) or (total_study_minutes > 0)

        if not has_data:
            return {
                "status": "success",
                "has_data": False,
                "primary_style": None,
                "style_description": None,
                "retention_rate": 0,
                "comprehension_speed": None,
                "peak_focus_hours": None,
                "cognitive_traits": {
                    "visual": 0,
                    "verbal": 0,
                    "analytical": 0,
                    "active_recall": 0,
                    "speed": 0
                },
                "strengths": [],
                "weaknesses": [],
                "ai_recommendations": []
            }

        avg_score = round(sum(q.get("percentage", 0) for q in quiz_records) / quizzes_taken) if quizzes_taken > 0 else 0

        # Determine Primary Style purely from actual user activity
        if quizzes_taken >= 2 and avg_score >= 70:
            primary_style = "Active Recall Specialist"
            style_desc = "Your activity shows high retention when taking practice quizzes and flashcard drills."
        elif docs_count >= 1:
            primary_style = "Visual & Textual Analyst"
            style_desc = "Your activity demonstrates strong synthesis of uploaded documents and written content."
        else:
            primary_style = "Interactive Conversational Tutor"
            style_desc = "Your learning pattern relies on dialogue and probing questions with your AI tutor."

        retention_rate = round(avg_score * 0.9) if quizzes_taken > 0 else 0
        speed_multiplier = f"{round(max(1.0, total_study_minutes / 30), 1)}x Study Speed" if total_study_minutes > 0 else "Normal Pace"
        peak_focus = "Active Study Sessions Recorded" if total_study_minutes > 0 else "N/A"

        cognitive_traits = {
            "visual": min(100, docs_count * 25),
            "verbal": min(100, docs_count * 15 + quizzes_taken * 10),
            "analytical": min(100, quizzes_taken * 25),
            "active_recall": avg_score,
            "speed": min(100, min(100, total_study_minutes * 2))
        }

        # Build dynamic strengths and weaknesses strictly from real data
        strengths = []
        if docs_count > 0:
            strengths.append(f"{docs_count} Document(s) Indexed in Knowledge Base")
        if quizzes_taken > 0:
            strengths.append(f"{quizzes_taken} Quiz Attempts Completed")
        if avg_score >= 75:
            strengths.append(f"High MCQ Accuracy ({avg_score}%)")

        weaknesses = []
        if quizzes_taken == 0:
            weaknesses.append("No practice quizzes completed yet")
        elif avg_score < 70:
            weaknesses.append(f"Quiz Accuracy needs improvement ({avg_score}%)")
        if docs_count == 0:
            weaknesses.append("No study documents uploaded yet")

        # Generate AI Recommendations dynamically using Groq LLM ONLY if user has uploaded docs/quizzes
        ai_recommendations = []
        if groq_client and docs_count > 0:
            try:
                doc_names = ", ".join([d.get("original_filename", "") for d in user_docs[:3]])
                prompt = f"""Student profile: Learning Style '{primary_style}', Quiz Accuracy {avg_score}%, Documents: {doc_names}.
Provide 3 concise, highly actionable study tips (1 sentence each) in strict JSON format:
[
  {{"title": "Tip Title", "action": "Action sentence.", "type": "quiz"}},
  {{"title": "Tip Title", "action": "Action sentence.", "type": "flashcards"}},
  {{"title": "Tip Title", "action": "Action sentence.", "type": "study-plan"}}
]"""
                llm_res = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3
                )
                parsed_tips = extract_json(llm_res.choices[0].message.content)
                if isinstance(parsed_tips, list):
                    ai_recommendations = parsed_tips[:3]
            except Exception as llm_err:
                logger.warning(f"Failed to generate LLM Learning DNA tips: {llm_err}")

        return {
            "status": "success",
            "has_data": True,
            "primary_style": primary_style,
            "style_description": style_desc,
            "retention_rate": retention_rate,
            "comprehension_speed": speed_multiplier,
            "peak_focus_hours": peak_focus,
            "cognitive_traits": cognitive_traits,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "ai_recommendations": ai_recommendations
        }
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to fetch Learning DNA data.")

# --- 🕸️ 3. INTERACTIVE KNOWLEDGE MAP ENDPOINT ---
@app.get("/api/knowledge-map")
async def get_knowledge_map(user_id: str = Depends(get_current_user)):
    try:
        user_docs = await db.user_documents.find({"user_id": user_id}).to_list(length=50)
        quiz_records = await db.quiz_attempts.find({"user_id": user_id}).to_list(length=50)

        # Strictly return empty state if user has uploaded no documents
        if not user_docs:
            return {
                "status": "success",
                "has_data": False,
                "nodes": [],
                "edges": [],
                "total_nodes": 0,
                "completed_count": 0,
                "in_progress_count": 0,
                "pending_count": 0
            }

        nodes = []
        edges = []

        for index, doc in enumerate(user_docs):
            doc_name = doc.get("doc_name")
            filename = doc.get("original_filename", doc_name)
            chunks = doc.get("chunks_count", 0)

            # Match actual quiz attempts for this doc
            doc_quizzes = [q for q in quiz_records if q.get("doc_name") == doc_name]
            if doc_quizzes:
                mastery = round(sum(q.get("percentage", 0) for q in doc_quizzes) / len(doc_quizzes))
                status = "completed" if mastery >= 75 else "in-progress"
            elif quiz_records:
                mastery = 50
                status = "in-progress"
            else:
                mastery = 0
                status = "pending"

            node_id = f"node_doc_{index+1}"
            nodes.append({
                "id": node_id,
                "label": filename.replace(".pdf", "").replace(".docx", "").replace("_", " ").title(),
                "category": "Document Core",
                "status": status,
                "mastery": mastery,
                "description": f"Uploaded document '{filename}' with {chunks} indexed context chunks.",
                "doc_name": doc_name,
                "subtopics": [f"{chunks} Indexed Chunks", "Semantic RAG Search", "Automated MCQs"]
            })

            # Connect nodes sequentially
            if index > 0:
                edges.append({
                    "source": f"node_doc_{index}",
                    "target": node_id,
                    "label": "Prerequisite"
                })

        return {
            "status": "success",
            "has_data": True,
            "nodes": nodes,
            "edges": edges,
            "total_nodes": len(nodes),
            "completed_count": sum(1 for n in nodes if n["status"] == "completed"),
            "in_progress_count": sum(1 for n in nodes if n["status"] == "in-progress"),
            "pending_count": sum(1 for n in nodes if n["status"] == "pending")
        }
    except Exception as e:
        logger.exception(e)
        raise HTTPException(status_code=500, detail="Failed to fetch Knowledge Map graph.")

# --- HEALTH CHECK ---
@app.get("/health")
async def health():
    return {"status": "ok"}

@app.get("/")
async def root():
    return {"status": "online", "message": "LearnNova AI Enterprise Production Backend is active."}